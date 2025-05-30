from flask import make_response, session, url_for, Blueprint, render_template, request, flash, redirect, current_app, jsonify
from .modules.user_mgt_module.staff import Staff
from .modules.user_mgt_module.admin import Admin
from .modules.login_register_module import Authorization
from .database.dbhelper import Databasehelper
from werkzeug.utils import secure_filename
import os
from datetime import datetime
import csv
import io
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.units import inch


admin = Blueprint('admin', __name__, template_folder='templates/admin',url_prefix='/admin')

auth = Authorization()

admin_account = None
db = Databasehelper()

# Allowed extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
ALLOWED_FILE_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx'}

# Function to check if the file extension is valid
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
def allowed_file_extension(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_FILE_EXTENSIONS

def save_image(image):
     # ✅ Attempt to save image
    file = image
    filename = ''

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename) # Secure file name
        upload_folder = os.path.abspath(os.path.join(current_app.root_path, "static/images/laboratories"))
        filepath = os.path.join(upload_folder, filename)

        if not os.path.isdir(upload_folder):
            print(f"⚠️ ERROR: Directory {upload_folder} does NOT exist!")
            flash("Upload folder missing. Contact admin.", "danger")
            return redirect(url_for('admin.admin_announcement'))
        
        print(f"Saving file to: {filepath}")  # Debugging line
        file.save(filepath.replace("\\", "/"))  # Fix Windows backslash issue

        if os.path.exists(filepath):  # Verify if the file was actually saved
            print("✅ File successfully saved!")
        
    return filename


@admin.after_request
def after_request(response):
    response.headers['Cache-Control'] = 'no-cache,no-store,must-revalidate'
    return response 

@admin.route('/login')
def adminlogin():
    return render_template('adminlogin.html')

@admin.route('/dashboard')
def dashboard():
    """Admin Dashboard"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'admindashboard.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account
                               )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/adminleaderboards')
def adminleaderboards():
    """Admin Dashboard"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'admin_leaderboards.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                               )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))   

@admin.route('/statistics')
def statistics():
    """Admin Dashboard"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'admin_statistics.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account
                               )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
       
    

@admin.route('/admin_announcements')
def admin_announcements():
    """Announcements"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        response = admin_account.get_announcements_for_students()

        if response['success']:
            return render_template(
                                    'admin_announcements.html',
                                    user_in_login_page=True, 
                                    action='Logout',
                                    admin=admin_account,
                                    )
        else:
            flash('An error occured', 'error')
            return redirect(url_for('admin.dashboard'))

    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    
@admin.route('/users')
def users():
    """Users"""
    global admin_account
    dash_search = request.args.get('dash_search', None)  # Get search query from URL parameters
    
    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
            'users.html',
            user_in_login_page=True,
            action='Logout',
            admin=admin_account,
            dash_search=dash_search  # Pass the search term to the template
        )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))




@admin.route('/managelab')
def managelab():
    """Manage Laboratories"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'managelab.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    
@admin.route('/resources')
def resources():
    """Manage Laboratories"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'lab_resources_materials.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/adminrecords')
def adminrecords():
    """Records"""
    global admin_account
    if not session['admin'] == None:
        
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)
        

        return render_template(
                                'adminrecords.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    

@admin.route('/admin_reservation_requests')
def admin_reservation_requests():
    """Reservation Requests"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'reservation_requests.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/adminusermgt')
def adminusermgt():
    """Admin User Management"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)
            
        staffs = admin_account.retrieve_all_staff()
        

        return render_template(
                                'adminusermgt.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                staffs=staffs,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))


@admin.route('/feedbacks')
def admin_feedbacks():
    """ Admin Feedbacks Management"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'admin_feedbacks.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/logs')
def logs():
    """Render Logs HTML"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        return render_template(
                                'admin_logs.html',
                                user_in_login_page=True, 
                                action='Logout',
                                admin=admin_account,
                                )


@admin.route('/adminlogout')
def adminlogout():
    session['admin'] = None
    flash('Admin logged out', 'error')
    return redirect(url_for('admin.adminlogin'))


@admin.route('/api/users')
def api_users():
    """API for fetching of users"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        users = admin_account.retrieve_all_students()

        # Convert objects to dictionaries
        users_list = []
        for user in users:
            users_list.append(
                {
                "idno": user['idno'],
                "firstname": user['firstname'],
                "middlename": user['middlename'],
                "lastname": user['lastname'],
                "course": user['course'],
                "year": user['year'],
                "email": user['email'],
                "session": user['session'],
                "status": user['status']
                }
            )

        return jsonify(users_list)  # Send JSON response
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    

@admin.route('/api/users/search', methods=['GET'])
def search_users():
    """API for fetching of filtered users"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        query = request.args.get('q','').strip()

        users = admin_account.retrieve_all_students()

        filtered_users = [
            {
                "idno": user['idno'],
                "firstname": user['firstname'],
                "middlename": user['middlename'],
                "lastname": user['lastname'],
                "course": user['course'],
                "year": user['year'],
                "email": user['email'],
                "session": user['session'],
                "status": user['status']
            }
            for user in users if (
                query.lower() in user['idno'].lower() or  # ✅ FIXED: user['idno'] (Dictionary)
                query.lower() in user['firstname'].lower() or
                query.lower() in user['middlename'].lower() or
                query.lower() in user['lastname'].lower() or
                query.lower() in user['email'].lower()
            )
        ]

        return jsonify(filtered_users)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/users/filter-status', methods=['GET'])
def filter_status():
    """API for fetching of filtered users"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        query = request.args.get('q', '')

        users = admin_account.retrieve_all_students()

        if not query == 'all':
            filtered_users_status = [
                {
                    "idno": user['idno'],
                    "firstname": user['firstname'],
                    "middlename": user['middlename'],
                    "lastname": user['lastname'],
                    "course": user['course'],
                    "year": user['year'],
                    "email": user['email'],
                    "session": user['session'],
                    "status": user['status']
                }
                for user in users if (
                    query.lower() in user['status'].lower()  # ✅ FIXED: user['idno'] (Dictionary)
                )
            ]

            return jsonify(filtered_users_status)
        
        # Convert objects to dictionaries
        users_list = []
        for user in users:
            users_list.append(
                {
                "idno": user['idno'],
                "firstname": user['firstname'],
                "middlename": user['middlename'],
                "lastname": user['lastname'],
                "course": user['course'],
                "year": user['year'],
                "email": user['email'],
                "session": user['session'],
                "status": user['status']
                }
            )

        return jsonify(users_list)
    
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/users/<idno>/logoff')
def logoff(idno):
    """API to Log-off student from laboratory"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        response = admin_account.logoff_student(idno=idno, staff_idno=admin_account.idno)

        if response['success']:
            return jsonify({'success': True, 'message': 'Student logged off successfully!'})
        else:
            return jsonify({'success': False, 'message': response['message']}), 400
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/users/currentsitin')
def current_sitin():
    """ API for retrieving all students currently in-lab """
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        users = admin_account.retrieve_all_current_sitins()

        json_format_users = [
            {
                "idno": user['idno'],
                "firstname": user['firstname'],
                "middlename": user['middlename'],
                "lastname": user['lastname'],
                "course": user['course'],
                "year": user['year'],
                "email": user['email'],
                "session": user['session'],
                "status": user['status'],
            }
            for user in users
        ]

        return jsonify(json_format_users)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))



@admin.route('/api/sitinrecords')
def sitinrecords():
    """API to retrieve sitin records"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        users = admin_account.retrieve_all_sitinrecords()

        # Convert objects to dictionaries
        users_list = []
        for user in users['data']:
            users_list.append({
                "record_id": user["record_id"],
                "idno": user["idno"],
                "lab_name": user["lab_name"],
                "sitin_in": user["sitin_in"],
                "sitin_out": user["sitin_out"],
                "staff_idno": user["staff_idno"],
                "logged_off_by": user["logged_off_by"],
                "status": user["status"],
                "reason": user["reason"],
            })

        return jsonify(users_list)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    
@admin.route("/api/reservations")
def get_pending_reservations():
    """API to retrieve all reservations"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = session.get('admin')
            admin_account = Admin(**admin_account)

        reservations = admin_account.retrieve_all_pending_reservations()

        # Convert objects to dictionaries
        reservations_list = []
        for reservation in reservations:
            reservations_list.append({
                "reservation_id": reservation["reservation_id"],
                "idno": reservation["idno"],
                "fullname" : reservation["fullname"],
                "lab_name": reservation["lab_name"],
                "reserve_date": reservation["reserve_date"],
                "request_date": reservation["request_date"],
                "status": reservation["status"],
                "computer": reservation["computer"],
                "purpose": reservation["purpose"],

            })

        return jsonify(reservations_list)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
 
@admin.route('/admin_editstudent', methods=['POST'])
def admin_editstudent():
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = Admin(**session.get('admin'))

        form_data = {
            "idno": request.form["idno"],
            "firstname": request.form["firstname"],
            "middlename": request.form["middlename"],
            "lastname": request.form["lastname"],
            "email": request.form["email"],
            "course": request.form["course"],
            "year": request.form["year"],
            "password": request.form["password"],
            "confirmpassword": request.form["confirmpassword"]
        }


        if form_data['password'] != form_data['confirmpassword']:
            flash('Password do not match', 'error')
            return redirect(url_for('admin.users'))

        del form_data['confirmpassword']

        response = admin_account.update(**{k:v for k,v in form_data.items() if v != ''})


        if response['success']:
            flash(response['message'], 'success')
            return redirect(url_for('admin.users'))
        else:
            flash(response['error'], 'error')
            return redirect(url_for('admin.users'))
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/adminlogin', methods=['POST'])
def loginadmin():
    global admin_account
    idno:str = request.form['idno']
    password:str = request.form['password']

    response = auth.user_account_exist_and_correct_credentials(
                                                                idno=idno,
                                                                password=password,
                                                                url='admin',
                                                                ) 
    if response['success']:
        del response['success']
        flash('Login successful!', 'success')
        session['admin'] = response
        admin_account = Admin(**response)
        return redirect(url_for('admin.dashboard'))
    else:
        flash(response['error'], 'error')
        return redirect(url_for('admin.adminlogin'))


@admin.route('/announce', methods=['POST'])
def announce():
    """Handle Announcements Post REQUESTS"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = Admin(**session.get('admin'))

        file = request.files['image']
        
        filename = save_image(image=file)

        announcement_data = {
            "post_title": request.form['title'],
            "post_description": request.form['description'],
            "image": filename,
            "posted_by":admin_account.idno,
            "date_posted": datetime.now(),
        }


        response = admin_account.add_announcement(**{k:v for k,v in announcement_data.items() if not v is None})
        if response['success']:
            flash('Announcement added', 'success')
            return redirect(url_for('admin.admin_announcements'))
        else:
            flash(response['error'], 'error')
            return redirect(url_for('admin.admin_announcements'))
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    

@admin.route('/api/announcements')
def get_announcements():
    """FETCH ALL ANNOUNCEMENTS"""
    global admin_account
    if not session.get('admin'):
        return jsonify({
            'success': False,
            'message': 'Unauthorized access is prohibited.'
        }), 401  # Unauthorized status code

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    response = admin_account.get_announcements_for_students()

    if response['success']:
        announcements = []
        for announcement in response['data']:
            announcements.append({
                'post_id': announcement['post_id'],
                'post_title': announcement['post_title'],
                'post_description': announcement['post_description'],
                'image': announcement['image'],
                'posted_by': announcement['posted_by'],
                'date_posted': announcement['date_posted'],
            })

        return jsonify({
            'success': True,
            'data': announcements
        })
    else:
        return jsonify({
            'success': False,
            'message': response['message']
        })
    
@admin.route('/api/<post_id>/editannouncement')
def editAnnoucement(post_id):
    """Edit announcement for Admin"""
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = Admin(**session.get('admin'))

        announcement = admin_account.get_announcement(post_id=post_id)

        # In your Flask route
        if announcement['success'] and announcement['data']:
            # Assuming you only want the first announcement if multiple exist
            data = announcement['data'][0]
            return jsonify({
                'post_id': data['post_id'],
                'post_title': data['post_title'],
                'post_description': data['post_description'],
                'image': data['image'],
                'posted_by': data['posted_by'],
                'date_posted': data['date_posted'],
            })
        else:
            return jsonify({'success': False, 'message': announcement['message']}), 400
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    
@admin.route('/api/<post_id>/update', methods=['POST'])
def update_announcement_one(post_id):
    """Update an existing announcement"""
    global admin_account
    if not session.get('admin'):
        flash('Unauthorized Access', 'error')
        return redirect(url_for('admin.adminlogin'))
    
    try:
        # Initialize admin account
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        # Get form data using same pattern as working announce() route
        update_data = {
            "post_title": request.form['editTitle'],  # Match form field name
            "post_description": request.form['editDescription'],
            "posted_by": admin_account.idno,
            "date_posted": datetime.now(),
        }

        # Handle file upload like in announce()
        if 'editPhoto' in request.files:  # Match file input name
            file = request.files['editPhoto']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                upload_folder = os.path.abspath(os.path.join(
                    current_app.root_path, 
                    "static/images/announcements"
                ))
                
                # Create directory if missing
                os.makedirs(upload_folder, exist_ok=True)
                
                filepath = os.path.join(upload_folder, filename)
                file.save(filepath)
                update_data['image'] = filename

        # Update using same pattern as add_announcement
        response = admin_account.update_announcement(
            post_id=post_id,
            **{k: v for k, v in update_data.items() if v is not None}
        )
        
        if response['success']:
            return jsonify({'success': True, 'message': 'Updated'})
        else:
            return jsonify({'success': False, 'message': response['message']}), 400

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 50

    
@admin.route('/api/announcements/<post_id>', methods=['DELETE'])
def delete_announcement(post_id):
    """Delete an announcement"""
    global admin_account

    print(f"This is the announcement ID : {post_id}")

    # Check if the admin is logged in
    if not session.get('admin'):
        return jsonify({
            'success': False,
            'message': 'Unauthorized access is prohibited.'
        }), 401  # Unauthorized status code

    image_filename = admin_account.get_announcement(post_id=post_id)

    if image_filename['success'] and image_filename['data']:
        # Delete the associated image
        image_path = os.path.join(
            current_app.static_folder,
            'images',
            'announcements',
            image_filename['data'][0]['image']
        )
        if os.path.exists(image_path):
            os.remove(image_path)
            current_app.logger.info(f"Deleted announcement image: {image_path}")

    # Initialize admin_account if not already initialized
    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    # Delete the announcement
    response = admin_account.delete_announcement(post_id=post_id)
    print(f"Response {response}")
    # Return the response
    return jsonify(response)

@admin.route('/api/approve_reservation/<reservation_id>')
def approve_reserve(reservation_id):
    """API to approve reservation"""
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        response = admin_account.approve_reservation(reservation_id=reservation_id,admin=admin_account)

        if response['success']:
            return jsonify({'success': True, 'message': 'Reservation approved successfully!'})
        else:
            return jsonify({'success': False, 'message': response['message']}), 400
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/decline_reservation/<reservation_id>')
def decline_reserve(reservation_id):
    """API to decline reservation"""
    global admin_account

    try:
        if session.get('admin') is None:
            return jsonify({
                'success': False,
                'message': 'Unauthorized access is prohibited'
            }), 401

        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        response = admin_account.decline_reservation(
            reservation_id=reservation_id,
            admin=admin_account
        )

        if response['success']:
            return jsonify({
                'success': True,
                'message': 'Reservation declined successfully!'
            })
        else:
            return jsonify({
                'success': False,
                'message': response['message']
            }), 400

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Server error: {str(e)}'
        }), 500

@admin.route('/sitin_student', methods=['POST'])
def sitin_student():
    """Sit In Student"""
    global admin_account
    if  session.get('admin'):

        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        idno = request.form['idno']
        lab_id = request.form['lab_id']
        reason = request.form['reason']

        if lab_id == "#" or reason == "#":
            return jsonify({'success': False, 'error': 'Please select a valid lab and purpose'}), 400

        response = admin_account.sit_in_student(idno=idno, lab_id=lab_id, reason=reason, staff_idno=admin_account.idno)

        if response['success']:
            return jsonify({'success': True, 'message': 'Student successfully sat in'})

        return jsonify({'success': False, 'error': response['error']}), 400
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))



@admin.route('/addstaff', methods=['POST'])
def addstaff():
    global admin_account
    if not session['admin'] == None:
        if admin_account == None:
            admin_account = Admin(**session.get('admin'))

        
        staff_data = {
                    'idno': request.form.get('idno'),
                    'firstname': request.form.get('firstname'),
                    'middlename': request.form.get('middlename'),
                    'lastname': request.form.get('lastname'),
                    'email': request.form.get('email'),
                    'course': request.form.get('course', ''),  # Default to empty string if not provided
                    'year': request.form.get('year', ''),  
                    'image': None,    # Set to None but will have new value later
                    }   

        password = request.form['password']
        confirmpassword = request.form['confirmpassword']

        if password != confirmpassword:
            flash('Password do not match', 'error')
            return redirect(url_for('admin.adminusermgt'))
        
        


        staff = Staff(**{k: v for k, v in staff_data.items() if v != ''})
        print(staff.__dict__)

        response = admin_account.add_staff(**staff.__dict__,password=password)

        if response['success']:
            flash('Successfully added','success')
            return redirect(url_for('admin.adminusermgt'))
        else:
             flash(response['error'],'error')
             return redirect(url_for('admin.adminusermgt'))
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    

@admin.route('/api/feedbacks')
def fetch_feedbacks():
    """ API for fetching feedbacks"""
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        feedbacks = admin_account.retrieve_all_feedbacks()

        json_formatted_data = [
            {
                "feedback_id": feedback['feedback_id'],
                "idno": feedback['idno'],
                "lab_name": feedback['lab_name'],
                "feedback": feedback['feedback'],
                "submitted_on": feedback['submitted_on'],
                "reason": feedback['reason']
            }
            for feedback in feedbacks['data']
        ]

        return jsonify(
            json_formatted_data
        )

    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/feedback/<feedback_id>')
def get_feedback(feedback_id):
    """ API for fetching single feedback by ID """
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        feedbacks = admin_account.retrieve_one_feedback(feedback_id=feedback_id)

        # Debug print to check the structure
        print(f"Raw feedback data: {feedbacks}")
        
        if not feedbacks.get('data') or not feedbacks['data']:
            return jsonify({"error": "Feedback not found"}), 404

        # Get the first feedback directly
        feedback = feedbacks['data'][0]
        print(f"Lab name: {feedback['lab_name']}")  # Debug print

        # Create response object directly
        response_data = {
            "feedback_id": feedback['feedback_id'],
            "idno": feedback['idno'],
            "lastname": feedback['lastname'],
            "firstname": feedback['firstname'],
            "lab_name": feedback['lab_name'],
            "feedback": feedback['feedback'],
            "submitted_on": feedback['submitted_on'],
            "reason": feedback['reason']
        }

        return jsonify(response_data)

    else:
        return jsonify({"error": "Unauthorized"}), 401


@admin.route('/api/records-by-lab/<lab_name>')
def fetch_records_by_lab(lab_name):
    """Fetch records by lab"""
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        records = admin_account.retrieve_sitinrecord_by_lab(lab_name=lab_name)


        json_formatted_data = [
            {
                "record_id": record['record_id'],
                "idno": record['idno'],
                "lab_name": record['lab_name'],
                "sitin_in": record['sitin_in'],
                "sitin_out": record['sitin_out'],
                "staff_idno": record['staff_idno'],
                "logged_off_by": record['logged_off_by'],
                "status": record['status'],
                "reason": record['reason']
            }
            for record in records['data']
        ]
        return jsonify(json_formatted_data)

    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/records-by-purpose/<purpose>')
def fetch_records_by_purpose(purpose):
    """Fetch records by purpose"""
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        records = admin_account.retrieve_sitinrecord_by_purpose(reason=purpose)

        json_formatted_data = [
            {
                "record_id": record['record_id'],
                "idno": record['idno'],
                "lab_name": record['lab_name'],
                "sitin_in": record['sitin_in'],
                "sitin_out": record['sitin_out'],
                "staff_idno": record['staff_idno'],
                "logged_off_by": record['logged_off_by'],
                "status": record['status'],
                "reason": record['reason']
            }
            for record in records['data']
        ]
        return jsonify(json_formatted_data)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))
    
@admin.route('/api/records-by-lab/<lab_name>/<purpose>')
def fetch_records_by_labpurpose(lab_name, purpose):
    """Fetch records by lab"""
    global admin_account

    if session.get('admin') is not None:
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        records = admin_account.retrieve_sitinrecord_by_lab_and_purpose(lab_name=lab_name,reason=purpose)


        json_formatted_data = [
            {
                "record_id": record['record_id'],
                "idno": record['idno'],
                "lab_name": record['lab_name'],
                "sitin_in": record['sitin_in'],
                "sitin_out": record['sitin_out'],
                "staff_idno": record['staff_idno'],
                "logged_off_by": record['logged_off_by'],
                "status": record['status'],
                "reason": record['reason']
            }
            for record in records['data']
        ]
        return jsonify(json_formatted_data)

    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))


@admin.route('/api/export-records/csv/<lab_name>')
@admin.route('/api/export-records/csv')
def export_records_csv(lab_name=None):
    """Export records to CSV"""
    global admin_account
    if session.get('admin') is None:
        return redirect(url_for('admin.adminlogin'))

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    # Get filtered records
    records = admin_account.retrieve_sitinrecord_by_lab(lab_id=lab_name) if lab_name else admin_account.retrieve_all_sitinrecords()

    # Create CSV in memory
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header
    writer.writerow(['Record ID', 'Student ID', 'Lab', 'Check-in', 'Check-out', 
                    'Staff ID', 'Logged Off By', 'Status', 'Reason'])
    
    # Write data
    for record in records:
        writer.writerow([
            record['record_id'],
            record['idno'],
            record['lab_name'],
            record['sitin_in'],
            record['sitin_out'] or '',
            record['staff_idno'] or '',
            record['logged_off_by'] or '',
            record['status'],
            record['reason'] or ''
        ])
    
    # Prepare response
    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=sit_in_records.csv'
    response.headers['Content-type'] = 'text/csv'
    return response

@admin.route('/api/export-records/excel/<lab_name>/<purpose>')
def export_records_excel(lab_name=None,purpose=None):
    """Export records to Excel"""
    global admin_account
    if session.get('admin') is None:
        return redirect(url_for('admin.adminlogin'))

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    if lab_name and lab_name != 'all' and purpose and purpose != 'all':
    # Filter by both lab_name and purpose
        print("IF CONDITION")
        records = admin_account.retrieve_sitinrecord_by_lab_and_purpose(lab_name=lab_name, reason=purpose)
    elif lab_name and lab_name != 'all':
        print("1st elif CONDITION")
        # Filter by lab_name only
        records = admin_account.retrieve_sitinrecord_by_lab(lab_name=lab_name)
    elif purpose and purpose != 'all':
        print("2ND elif CONDITION")
        # Filter by purpose only
        records = admin_account.retrieve_sitinrecord_by_purpose(reason=purpose)
        print(records)
    else:
        # No filters applied, retrieve all records
        print("ELSE CONDITION")
        records = admin_account.retrieve_all_sitinrecords()
#    # Fetch records
#     records = (admin_account.retrieve_sitinrecord_by_lab(lab_name=lab_name)
#                if lab_name and lab_name != 'all' else admin_account.retrieve_all_sitinrecords())

    # Create DataFrame
    df = pd.DataFrame([{
        'Record ID': r['record_id'],
        'Student ID': r['idno'],
        'Lab': r['lab_name'],
        'Check-in': r['sitin_in'],
        'Check-out': r['sitin_out'] or '',
        'Staff ID': r['staff_idno'] or '',
        'Logged Off By': r['logged_off_by'] or '',
        'Status': r['status'],
        'Reason': r['reason'] or ''
    } for r in records['data']])  # Assuming all keys exist


    # Create Excel in memory
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Sit-In Records')
        writer.close()  # Correct method

    output.seek(0)

    # Prepare response
    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = f'attachment; filename=sit_in_records_lab-{lab_name}_purpose-{purpose}.xlsx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    return response

@admin.route('/api/export-records/pdf/<lab_name>/<purpose>')
def export_records_pdf(lab_name=None, purpose=None):
    """Generate a well-formatted PDF for Sit-In Records."""
    global admin_account
    if session.get('admin') is None:
        return redirect(url_for('admin.adminlogin'))

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))
    
    if lab_name and lab_name != 'all' and purpose and purpose != 'all':
    # Filter by both lab_name and purpose
        print("IF CONDITION")
        records = admin_account.retrieve_sitinrecord_by_lab_and_purpose(lab_name=lab_name, reason=purpose)
    elif lab_name and lab_name != 'all':
        print("1st elif CONDITION")
        # Filter by lab_name only
        records = admin_account.retrieve_sitinrecord_by_lab(lab_name=lab_name)
    elif purpose and purpose != 'all':
        print("2ND elif CONDITION")
        # Filter by purpose only
        records = admin_account.retrieve_sitinrecord_by_purpose(reason=purpose)
        print(records)
    else:
        # No filters applied, retrieve all records
        print("ELSE CONDITION")
        records = admin_account.retrieve_all_sitinrecords()

    filename = f"sit_in_records_lab_{lab_name}.pdf" if lab_name else "sit_in_records_all.pdf"
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=landscape(letter), leftMargin=30, rightMargin=30, topMargin=50, bottomMargin=30)
    elements = []

    
    # University Header (only appears on the first page)
    university_header = [
        Paragraph("<b>University of Cebu Main Campus</b>", 
                  ParagraphStyle(name='Center', alignment=1, fontSize=16, fontName='Helvetica-Bold',leading=20,spaceAfter=7)),
        Paragraph("College of Computer Studies", 
                  ParagraphStyle(name='Center', alignment=1, fontSize=14,leading=20,spaceAfter=7)),
        Paragraph(f"<b>Sit-In Records - {lab_name+ '-' + purpose if lab_name else 'All Labs'}</b>", 
                  ParagraphStyle(name='Center', alignment=1, fontSize=14, fontName='Helvetica-Bold',leading=20,spaceAfter=7)),
        Spacer(1, 0.5 * inch)  # Add space after header
    ]
 
    # Table Headers
    data = [['Record ID', 'Student ID', 'Lab', 'Check-in', 'Check-out', 'Status', 'Reason']]
    
    # Add Data Rows
    for record in records['data']:
        data.append([
            str(record['record_id']),
            str(record['idno']),
            record['lab_name'],
            format_date(record['sitin_in']),
            format_date(record['sitin_out']),
            record['status'],
            record['reason'] or ''
        ])
    
    # Create Table
    table = Table(data, colWidths=[70, 80, 60, 120, 120, 70, 140])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightskyblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10)
    ]))
    
    content = elements + university_header + [table]
    doc.build(content)
    output.seek(0)

    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = f'attachment; filename={filename}'
    response.headers['Content-type'] = 'application/pdf'
    return response

def format_date(datetime_str):
    """Format datetime string for display, handling milliseconds"""
    if not datetime_str:
        return ''
    
    try:
        # Try parsing with milliseconds first
        try:
            dt = datetime.strptime(datetime_str, '%Y-%m-%d %H:%M:%S.%f')
        except ValueError:
            # Fall back to parsing without milliseconds
            dt = datetime.strptime(datetime_str, '%Y-%m-%d %H:%M:%S')
        return dt.strftime('%m/%d/%Y %I:%M %p')
    except ValueError as e:
        print(f"Error formatting date: {datetime_str} - {str(e)}")
        return datetime_str  # Return original string if formatting fails

@admin.route('/api/export-records/word/<lab_name>/<purpose>')
def export_records_word(lab_name=None, purpose=None):
    """Export records to Word"""
    global admin_account
    if session.get('admin') is None:
        return redirect(url_for('admin.adminlogin'))

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))
    
    if lab_name and lab_name != 'all' and purpose and purpose != 'all':
    # Filter by both lab_name and purpose
        print("IF CONDITION")
        records = admin_account.retrieve_sitinrecord_by_lab_and_purpose(lab_name=lab_name, reason=purpose)
    elif lab_name and lab_name != 'all':
        print("1st elif CONDITION")
        # Filter by lab_name only
        records = admin_account.retrieve_sitinrecord_by_lab(lab_name=lab_name)
    elif purpose and purpose != 'all':
        print("2ND elif CONDITION")
        # Filter by purpose only
        records = admin_account.retrieve_sitinrecord_by_purpose(reason=purpose)
        print(records)
    else:
        # No filters applied, retrieve all records
        print("ELSE CONDITION")
        records = admin_account.retrieve_all_sitinrecords()

    # Create Word document
    document = Document()

     # Configure the first section to have a different first page header
    section = document.sections[0]
    section.different_first_page_header_footer = True  # This ensures the header appears only on the first page
    
    # Add Header (only appears on the first page due to different_first_page_header_footer)
    header = section.first_page_header
    header_para = header.paragraphs[0]
    
    # University Name - make it bold and centered
    university = header_para.add_run("University of Cebu Main Campus")
    university.bold = True
    university.font.size = Pt(14)
    
    # College Name
    header_para.add_run("\nCollege of Computer Studies").font.size = Pt(12)
    
    # Report Title
    header_para.add_run("\nLaboratory Sit-in Monitoring").bold = True
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add some space after header
    document.add_paragraph("\n")
    
    # Create table
    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Add headers
    headers = table.rows[0].cells
    headers[0].text = 'Record ID'
    headers[1].text = 'Student ID'
    headers[2].text = 'Lab'
    headers[3].text = 'Check-in'
    headers[4].text = 'Check-out'
    headers[5].text = 'Status'
    headers[6].text = 'Reason'
    
    # Add data
    for record in records['data']:
        row = table.add_row().cells
        row[0].text = str(record['record_id'])
        row[1].text = str(record['idno'])
        row[2].text = record['lab_name']
        row[3].text = format_date(record['sitin_in'])
        row[4].text = format_date(record['sitin_out']) or ''
        row[5].text = record['status']
        row[6].text = record['reason'] or ''
    
    # Save to memory
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    # Prepare response
    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = f'attachment; filename=sit_in_records_lab-{lab_name}_purpose-{purpose}.docx'
    response.headers['Content-type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response




@admin.route('/api/add-laboratory', methods=['POST'])
def add_laboratory():
    """API to add a new laboratory"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        try:
            # ✅ Get form data
            lab_name = request.form.get("lab_name")
            lab_description = request.form.get("lab_description")
            vacant_time = request.form.get("vacant_time")
            day_sched = request.form.get("day_sched")
            image = request.files.get("image")  # Get file
            slots = request.form.get("slots")

            # ✅ Check if all fields are provided
            if not all([lab_name, lab_description, vacant_time, day_sched, image]):
                return jsonify({'success': False, 'message': 'All fields are required'}), 400

            # ✅ Attempt to save image
            file = image
            filename = ''

            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename) # Secure file name
                upload_folder = os.path.abspath(os.path.join(current_app.root_path, "static/images/laboratories"))
                filepath = os.path.join(upload_folder, filename)

                if not os.path.isdir(upload_folder):
                    print(f"⚠️ ERROR: Directory {upload_folder} does NOT exist!")
                    flash("Upload folder missing. Contact admin.", "danger")
                    return redirect(url_for('admin.admin_announcement'))
                
                print(f"Saving file to: {filepath}")  # Debugging line
                file.save(filepath.replace("\\", "/"))  # Fix Windows backslash issue

                if os.path.exists(filepath):  # Verify if the file was actually saved
                    print("✅ File successfully saved!")


            # ✅ Prepare data to send to `admin_account.add_lab`
            data = {
                "lab_name": lab_name,
                "lab_description": lab_description,
                "vacant_time": vacant_time,
                "day_sched": day_sched,
                "image": filename,
                "slots":slots,
                "admin_idno": admin_account.idno
            }

            response = admin_account.add_lab(**data)
            if response['success']:
                return jsonify({'success': True, 'message': 'Laboratory added successfully!'}), 200
            else:
                return jsonify({'success': False, 'message': response['error']}), 500

        except Exception as e:
            return jsonify({'success': False, 'message': str(e)}), 500

    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))


@admin.route('/api/get-laboratories')
def get_laboratories():
    """API to retrieve all laboratories"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        laboratories = admin_account.retrieve_all_labs()

        # Convert objects to dictionaries
        labs_list = []
        for lab in laboratories:
            labs_list.append({
                "lab_id": lab["lab_id"],
                "lab_name": lab["lab_name"],
                "lab_description": lab["lab_description"],
                "vacant_time": lab["vacant_time"],
                "slots": lab["slots"],
            })

        return jsonify(labs_list)
    else:
        flash('Unauthorized Access is Prohibited', 'error')
        return redirect(url_for('admin.adminlogin'))

@admin.route('/api/edit-laboratory/<lab_id>', methods=['POST'])
def edit_laboratory(lab_id):
    """API to edit laboratory details"""
    global admin_account

    if session.get('admin'):
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        try:
            # Get Form Data
            lab_name = request.form.get("lab_name")
            lab_description = request.form.get("lab_description")
            vacant_time = request.form.get("vacant_time")
            day_sched = request.form.get("day_sched")
            slots = request.form.get("slots")
            image = request.files.get("image")  # Get file

            filename = save_image(image)

            data = {
                'lab_id': lab_id,
                'lab_name':lab_name,
                'lab_description': lab_description,
                'vacant_time': vacant_time,
                'slots': slots,
                'day_sched': day_sched,
                'image':filename,
                'admin_idno': admin_account.idno,
            }
            response = admin_account.edit_lab(**{k:v for k,v in data.items() if not v == ''})
            if response['success']:
                return jsonify({'success': True, 'message': 'Laboratory edited successfully!'}), 200
            else:
                return jsonify({'success': False, 'message': response['error']}), 500

        except Exception as e:
            return jsonify({
                'success': False,
                'message': str(e)
            }),500

@admin.route('/api/delete-laboratory/<lab_id>', methods=['DELETE'])
def delete_laboratory(lab_id):
    """API to delete a laboratory"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        image_filename = admin_account.get_lab_details(lab_id=lab_id)

        if image_filename['success'] and image_filename['data']:
            # Delete the associated image
            image_path = os.path.join(
                current_app.static_folder,
                'images',
                'laboratories',
                image_filename['data'][0]['image']
            )
            if os.path.exists(image_path):
                os.remove(image_path)
                current_app.logger.info(f"Deleted laboratory image: {image_path}")
        # Call the delete method from your Admin class
        response = admin_account.delete_lab(lab_id=lab_id)
        
        if response['success']:
            return jsonify({
                'success': True,
                'message': 'Laboratory deleted successfully!'
            }), 200
        else:
            return jsonify({
                'success': False,
                'message': response.get('error', 'Failed to delete laboratory')
            }), 500
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

@admin.route('/api/labs/update-computer/<lab_id>', methods=['POST'])
def update_computer(lab_id):
    """Update Laboratory Computers Slot"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))
    if not request.is_json:
        return jsonify({'success': False, 'message': 'Missing JSON'}), 400
        
    computerStatusChanges = request.get_json()
    
    try:
        # Get the new number of slots from the request
        update_index = list(computerStatusChanges.keys())
        status = list(computerStatusChanges.values())
        print(update_index)
        print(status)

        if update_index is None:
            return jsonify({'success': False, 'message': 'Slots parameter is required'}), 400

        # Call the update method from your Admin class
        response = admin_account.update_computer_slots(lab_id=lab_id, update_index=update_index, status=status)

        if response['success']:
            return jsonify({'success': True, 'message': 'Laboratory computers updated successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500
    except Exception as e:
        print(f'Error in catch {str(e)}')
        return jsonify({'success': False, 'message': str(e)}), 500

@admin.route('/api/labs/find-lab/<lab_id>')
def find_lab(lab_id):
    """Find Laboratory"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))    

    try:
        # Call the find method from your Admin class
        response = admin_account.find_lab(lab_id=lab_id)

        print(f'Response data {response['data'][0]}')

        computers_list = list(response['data'][0]['computers'])

        data = {
            'computers':computers_list,
            'slots': response['data'][0]['slots'],
        }


        if response['success']:
            print('Success!')
            return jsonify({'success': True, 'message': 'Laboratory found successfully!', 'data': data}), 200
        else:
            print('Error!')
            return jsonify({'success': False, 'message': response['error']}), 500
        
    except Exception as e:
        print(f'Error in Catch {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 500


@admin.route('/api/get-lab-resources', methods=['GET'])
def get_lab_resources():
    """API to retrieve all lab resources"""
    global admin_account
    if not session.get('admin'):
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        resources = admin_account.get_lab_resources()  

        # Convert objects to dictionaries
        resources_list = [{
            "resources_id": resource["resources_id"],
            "resources_name": resource["resources_name"],
            "description": resource["description"],
            "resource_type": resource["resource_type"],
            "resources_path": resource["resources_path"],
            "status": resource["status"],
            "upload_date": resource["upload_date"],
        } 
        for resource in resources['data']
        ]

        return jsonify({'success': True, 'resources': resources_list}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

    
@admin.route('/api/add-lab-resource', methods=['POST'])
def add_lab_resources():
    """API to add a new lab resource"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        file = request.files.get("resource_file")

        if file and allowed_file_extension(file.filename):
            filename = secure_filename(file.filename)

            # Create upload folder if it doesn't exist
            upload_folder = os.path.abspath(os.path.join(current_app.root_path, "static/lab-resources"))
            os.makedirs(upload_folder, exist_ok=True)

            # Save file
            file_path = os.path.join(upload_folder, filename)
            file.save(file_path)

            # Store relative path to DB (e.g., static/uploads/filename.pdf)
            resource_path = os.path.join('static', 'uploads', filename)

            data = {
                'resources_name': request.form.get("resource_name"),
                'description': request.form.get("description"),
                'resource_type': request.form.get("resource_type"),
                'resources_path': filename,  # pass the path instead of file object
                'status': request.form.get("status"),
                'upload_date': datetime.now(),
            }

            response = admin_account.add_lab_resource(**data)

            if response['success']:
                return jsonify({'success': True, 'message': 'Laboratory resource added successfully!'}), 200
            else:
                return jsonify({'success': False, 'message': response['error']}), 500
        else:
            return jsonify({'success': False, 'message': 'Invalid or missing file type.'}), 400

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    

@admin.route('/api/delete-lab-resource/<resources_id>', methods=['DELETE'])
def delete_lab_resource(resources_id):
    """API to delete a lab resource"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:

        deleteResourceFile(resources_id=resources_id)

        # Call the delete method from your Admin class
        response = admin_account.delete_lab_resource(resources_id=resources_id)
        
        if response['success']:
            return jsonify({'success': True, 'message': 'Laboratory resource deleted successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@admin.route('/api/edit-lab-resource/<resources_id>', methods=['POST'])
def edit_lab_resource(resources_id):
    """API to edit a lab resource"""
    global admin_account 

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        print('Admin Object is none')
        admin_account = Admin(**session.get('admin'))
        print('Admin object no longer None')

    try:

        filename = ''
        file = request.files.get("edit-file")
        

        if file and allowed_file_extension(file.filename):

            deleteResourceFile(resources_id=request.form.get("edit-resources_id"))

            filename = secure_filename(file.filename)
             # Create upload folder if it doesn't exist
            upload_folder = os.path.abspath(os.path.join(current_app.root_path, "static/lab-resources"))
            os.makedirs(upload_folder, exist_ok=True)

            # Save file
            file_path = os.path.join(upload_folder, filename)
            file.save(file_path)
            
            
        data = {
            'resources_id': resources_id,
            'resources_name': request.form.get("edit-resources_name"),
            'resources_path': filename,  # pass the path instead of file object
            'description': request.form.get("edit-description"),
            'resource_type': request.form.get("edit-resource_type"),
            'status': request.form.get("edit-status"),
            'upload_date': datetime.now(),
        }

        response = admin_account.update_lab_resource(**{k:v for k,v in data.items() if not v == ''})  

        if response['success']:
            return jsonify({'success': True, 'message': 'Laboratory resource edited successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500

    except Exception as  e:
         return jsonify({'success': False, 'message': str(e)}), 500

@admin.route('/api/publish-unpublish', methods=['POST'])
def publish_unpublish():
    """API to publish/unpublish lab resources"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        resources_id = request.form.get('resources_id')
        status = request.form.get('edit-status')

        response = admin_account.publish_unpublish(resources_id=resources_id, status=status)

        if response['success']:
            return jsonify({'success': True, 'message': 'Status updated successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500
    

def deleteResourceFile(resources_id):
        """Delete a resource file from static folder"""
     # Get the resource details first
        resource_details = admin_account.get_lab_resource_by_id(resources_id=resources_id)

        if resource_details['success'] and resource_details['data']:
            # Delete the associated file
            file_path = os.path.join(
                current_app.static_folder,
                'lab-resources',
                resource_details['data'][0]['resources_path']
            )
            if os.path.exists(file_path):
                os.remove(file_path)
                current_app.logger.info(f"Deleted lab resource file: {file_path}")


@admin.route('/api/reset-one-session/<idno>', methods=['POST'])
def reset_one(idno):
    """API to reset a single session"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        response = admin_account.reset_one_session(idno=idno)

        if response['success']:
            return jsonify({'success': True, 'message': 'Session reset successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500

    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500


@admin.route('/api/reset-all-session', methods=['POST'])
def reset_all():
    """API to reset all sessions"""
    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    try:
        # Initialize admin if needed
        global admin_account
        if admin_account is None:
            admin_account = Admin(**session.get('admin'))

        # Reset sessions
        success = admin_account.reset_all_sessions()
        
        if success:
            return jsonify({
                'success': True,
                'message': 'All sessions reset successfully!'
            }), 200
        else:
            return jsonify({
                'success': False,
                'message': 'Failed to reset sessions (database error)'
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Server error: {str(e)}'
        }), 500

@admin.route('/api/reset-password', methods=['POST'])
def reset_student_password():
    """API to reset password"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        idno = request.form.get('idno')
        new_password = request.form.get('newPassword')

        response = admin_account.reset_password(idno=idno, password=new_password)
        if response['success']:
            return jsonify({'success': True, 'message': 'Password reset successfully!'}), 200
        else:
            return jsonify({'success': False, 'message': response['error']}), 500
    
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@admin.route('/api/adminleaderboards')
def get_leaderboards():
    """Fetch Leaderboards"""
    global admin_account
    if not session.get('admin'):
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    if admin is None:
        admin_account = Admin(**session.get('admin'))

    try:
        leaderboards = admin_account.retrieve_leaderboards() 


        leaderboards_list = [
            {
                'Name': leaderboard['Name'],
                'Program': leaderboard['Program'],
                'session_count' : leaderboard['session_count'],
                'rank': leaderboard['rank'],
                'idno':leaderboard['idno'],
                'image':leaderboard['image'],
                'points': leaderboard['points']
            }
            for leaderboard in leaderboards['data']
        ] 


        return jsonify({'success': True, 'leaderboards': leaderboards_list}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@admin.route('/api/add-points/<idno>/<points>', methods=['POST'])
def add_points(idno, points):
    """Fetch API to add points to a student to be logged out"""
    global admin_account

    if not session.get('admin'):
        return jsonify({
            'success':False,
            'error': 'Unauthorized'}), 401
    if not admin_account:
        admin_account = Admin(**session.get('admin'))
    try:
        response = admin_account.add_points(idno=idno, points=points)
        if response['success']:
            return jsonify({
                'success':True
            }),200
        else:
            return jsonify({
                'success':False,
                'message': f'Failed to add points. Error: {response['error']}'
                            })
    except Exception as e:
        return jsonify({
            'success':False,
            'error': f'An error occured: {str(e)}'
        })

@admin.route('/api/retrieve-logs')
def retrieve_logs():
    """FETCH API Retrieve logs"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        logs = admin_account.retrieve_logs()

        logs_list = [
            {
                'logs_id': log['logs_id'],
                'reservation_id': log['reservation_id'],
                'student_idno': log['student_idno'],
                'fullname': log['fullname'],
                'action_taken_on': log['action_taken_on'],
                'status': log['status'],
                'staff_idno': log['staff_idno'],
                'request_date': log['request_date'],
                'reserve_date': log['reserve_date'],
                'lab_name': log['lab_name'],
                'computer': log['computer'],
                'purpose': log['purpose']
            }
            for log in logs['data']
        ]

        return jsonify({'success': True, 'data': logs_list}), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@admin.route('/api/auto-sitin')
def auto_sitin():
    global admin_account
    
    if not session.get('admin'):
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
    if admin_account is None:
        admin_account = Admin(**session.get('admin'))
    
    result = admin_account.auto_sitin_reserved_students()
    status_code = 200 if result['success'] else 400
    return jsonify(result), status_code


@admin.route('/api/sitin-purpose-stats')
def sitin_purpose_stats():
    """Get sit-in purpose statistics"""
    global admin_account

    if not session.get('admin'):
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401

    if admin_account is None:
        admin_account = Admin(**session.get('admin'))

    try:
        result = admin_account.get_purpose_stats()
        return jsonify(result), 200
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


    